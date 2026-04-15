"""
Strategy pattern: abstract base + concrete parsers for each file type.
This is Extensible so we can add new parsers without touching existing code.
"""
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict
import chardet


@dataclass
class PageContent:
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    text: str
    pages: List[PageContent]
    metadata: Dict = field(default_factory=dict)


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> ExtractedDocument:
        pass


class PDFParser(DocumentParser):
    def parse(self, file_path: str) -> ExtractedDocument:
        import fitz
        doc = fitz.open(file_path)
        pages = []
        full_text = []
        for i, page in enumerate(doc):
            text = page.get_text()
            pages.append(PageContent(page_number=i + 1, text=text))
            full_text.append(text)
        doc.close()
        return ExtractedDocument(
            text="\n\n".join(full_text),
            pages=pages,
            metadata={"page_count": len(pages), "source": Path(file_path).name}
        )


class DOCXParser(DocumentParser):
    def parse(self, file_path: str) -> ExtractedDocument:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return ExtractedDocument(
            text=full_text,
            pages=[PageContent(page_number=1, text=full_text)],
            metadata={"paragraph_count": len(paragraphs), "source": Path(file_path).name}
        )


class TextParser(DocumentParser):
    def parse(self, file_path: str) -> ExtractedDocument:
        raw = Path(file_path).read_bytes()
        encoding = chardet.detect(raw)["encoding"] or "utf-8"
        text = raw.decode(encoding)
        return ExtractedDocument(
            text=text,
            pages=[PageContent(page_number=1, text=text)],
            metadata={"encoding": encoding, "source": Path(file_path).name}
        )


# Map file extension to correct parser
PARSER_MAP = {
    ".pdf": PDFParser,
    ".docx": DOCXParser,
    ".txt": TextParser,
    ".md": TextParser,
}

ALLOWED_EXTENSIONS = set(PARSER_MAP.keys())


def get_parser(file_path: str) -> DocumentParser:
    ext = Path(file_path).suffix.lower()
    if ext not in PARSER_MAP:
        raise ValueError(f"Unsupported file type: {ext}. Allowed: {ALLOWED_EXTENSIONS}")
    return PARSER_MAP[ext]()